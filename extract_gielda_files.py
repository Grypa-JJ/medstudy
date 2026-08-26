import sys
from pathlib import Path
import fitz  # PyMuPDF
import docx

def extract_pdf(path):
    doc = fitz.open(path)
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)

def extract_docx(path):
    d = docx.Document(path)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)

def extract(path):
    path = Path(path)
    if path.suffix.lower() == '.pdf':
        return extract_pdf(path)
    elif path.suffix.lower() == '.docx':
        return extract_docx(path)
    else:
        raise ValueError(f"Unsupported: {path}")

if __name__ == '__main__':
    base = Path('Rok 3 2025-2026') / 'GIEŁDY TEGOROCZNE 2025-2026🌐'
    out_dir = Path('scratch_gielda_extract')
    out_dir.mkdir(exist_ok=True)

    targets = []
    for sesja in ['SESJA LETNIA ☀️', 'SESJA ZIMOWA ❄️']:
        sdir = base / sesja
        if not sdir.exists():
            continue
        for sub in sdir.iterdir():
            if sub.is_dir() and sub.name.lower() in ('patologia', 'medycyna sądowa'):
                for f in sub.iterdir():
                    if f.is_file() and f.suffix.lower() in ('.pdf', '.docx'):
                        targets.append(f)

    for f in targets:
        try:
            text = extract(f)
        except Exception as e:
            text = f"[ERROR: {e}]"
        safe_name = f.parent.parent.name.replace(' ', '_').replace('☀️','L').replace('❄️','Z') + '__' + f.parent.name + '__' + f.stem
        safe_name = safe_name.encode('ascii', 'ignore').decode('ascii') or f.stem
        out_path = out_dir / (safe_name[:80] + '.txt')
        out_path.write_text(text, encoding='utf-8')
        print(f"{f} -> {out_path} ({len(text)} chars)")
