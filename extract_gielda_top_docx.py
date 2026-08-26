import docx, os

base = r"Giełdy tegoroczne"
out = r"C:\Users\Jakub\AppData\Local\Temp\claude\C--Users-Jakub-Desktop-Prod-projekt-w-budowie\5b271756-3497-41b2-84e6-8e1d1037c3aa\scratchpad\gielda_tegoroczne_top"

files = [
    (r"GIEŁDA 2025_2026.docx", "top_level_GIELDA.txt"),
    (r"Anatomia🫀-20260719T204335Z-1-001\Anatomia🫀\Praktyka\GIEŁDA 2025_2026.docx", "anatomia_praktyka.txt"),
    (r"Anatomia🫀-20260719T204335Z-1-001\Anatomia🫀\Teoria\Dokument bez tytułu.docx", "anatomia_teoria.txt"),
    (r"Angielski🇬🇧-20260719T204328Z-1-001\Angielski🇬🇧\GIEŁDA 2025_2026.docx", "angielski.txt"),
    (r"NASZE GIEŁDY-20260719T204242Z-1-001\NASZE GIEŁDY\GIEŁDA 2025_2026.docx", "nasze_gieldy.txt"),
    (r"Praktyka-20260719T204307Z-1-001\Praktyka\GIEŁDA 2025_2026_.docx", "praktyka.txt"),
    (r"Teoria-20260719T204304Z-1-001\Teoria\GIEŁDA 2025_2026.docx", "teoria.txt"),
]

for relpath, outname in files:
    path = os.path.join(base, relpath)
    try:
        d = docx.Document(path)
        text = "\n".join(p.text for p in d.paragraphs)
        # also tables
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        outpath = os.path.join(out, outname)
        with open(outpath, "w", encoding="utf-8") as f:
            f.write(text)
        print(relpath, "->", len(text), "chars,", len(d.paragraphs), "paragraphs")
    except Exception as e:
        print(relpath, "ERROR:", e)
